from pathlib import Path
from tempfile import NamedTemporaryFile
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "results" / "reports" / "analysis_importance_legacy"
CHART_PATH = OUT_DIR / "activity_risk_comparison.png"


DOCS = [
    {
        "path": OUT_DIR / "Feature_Importance_Business_Insights_EN.docx",
        "anchor": "Business interpretation: customers who do not hold this core banking product",
        "heading": "Supplementary Figure — Activity, Zero Balance, and ind_var30 Risk Comparison",
        "caption": (
            "Three independent analyses converge on the same ~20K customer segment: Q1 lowest activity, "
            "saldo_var30 <= 0, and ind_var30 == 0 are each about 2.2x more likely to be dissatisfied than average. "
            "This reinforces that low activity, zero/negative balance, and no ind_var30 product holding are "
            "different views of the same actionable high-risk customer profile."
        ),
    },
]


def insert_after(anchor_paragraph, *elements):
    parent = anchor_paragraph._p.getparent()
    index = parent.index(anchor_paragraph._p)
    for offset, element in enumerate(elements, start=1):
        parent.insert(index + offset, element)


def build_insert_block(doc: Document, heading: str, caption: str):
    heading_para = doc.add_paragraph()
    try:
        heading_para.style = doc.styles["Heading3"]
    except KeyError:
        try:
            heading_para.style = "Heading 2"
        except KeyError:
            pass
    heading_run = heading_para.add_run(heading)
    heading_run.bold = True
    heading_run.font.size = Pt(12)

    image_para = doc.add_paragraph()
    image_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = image_para.add_run()
    run.add_picture(str(CHART_PATH), width=Inches(6.3))

    caption_para = doc.add_paragraph()
    caption_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = caption_para.add_run(caption)
    run.font.size = Pt(9)
    run.font.italic = True

    return heading_para._p, image_para._p, caption_para._p


def update_docx(doc_info: dict) -> None:
    doc = Document(doc_info["path"])
    if any(doc_info["heading"] in paragraph.text for paragraph in doc.paragraphs):
        for paragraph in doc.paragraphs:
            if (
                "This chart compares Austin" in paragraph.text
                or "Three independent analyses converge" in paragraph.text
            ):
                paragraph.text = doc_info["caption"]
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.italic = True
        doc.save(doc_info["path"])
        replace_embedded_chart(doc_info["path"])
        return

    anchor = next(
        paragraph
        for paragraph in doc.paragraphs
        if doc_info["anchor"] in paragraph.text
    )
    elements = build_insert_block(doc, doc_info["heading"], doc_info["caption"])
    insert_after(anchor, *elements)
    doc.save(doc_info["path"])
    replace_embedded_chart(doc_info["path"])


def replace_embedded_chart(docx_path: Path) -> None:
    chart_bytes = CHART_PATH.read_bytes()
    with ZipFile(docx_path, "r") as zin, NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        tmp_path = Path(tmp.name)
        with ZipFile(tmp, "w", ZIP_DEFLATED) as zout:
            replaced = False
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.startswith("word/media/") and item.filename.endswith(".png"):
                    data = chart_bytes
                    replaced = True
                zout.writestr(item, data)
            if not replaced:
                raise RuntimeError(f"No embedded PNG found in {docx_path}")
    tmp_path.replace(docx_path)


if __name__ == "__main__":
    for doc_info in DOCS:
        update_docx(doc_info)
        print(doc_info["path"])
